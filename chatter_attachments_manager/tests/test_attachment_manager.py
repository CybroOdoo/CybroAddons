# -*- coding: utf-8 -*-
from odoo.tests import TransactionCase
from odoo.exceptions import ValidationError
from odoo.tools import mute_logger
import base64
import io
import pandas as pd
from docx import Document as DocxDocument
import logging

_logger = logging.getLogger(__name__)


class TestAttachmentManager(TransactionCase):

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        _logger.info("========== STARTING TestAttachmentManager SETUPClass ==========")

        # Create base test partner
        cls.partner = cls.env['res.partner'].create({
            'name': 'Test Attachment Partner',
            'email': 'attachment.partner@example.com',
        })
        _logger.info("Test partner created: ID %s", cls.partner.id)

    def test_01_tag_constraints_and_defaults(self):
        """Test attachment tag default color generation and unique constraints."""
        _logger.info("========== RUNNING test_01_tag_constraints_and_defaults ==========")
        # 1. Test default color generation
        tag_1 = self.env['ir.attachment.tag'].create({
            'name': 'Invoice Tag',
        })
        _logger.info("Created Tag 1: Name: %s, Color: %s", tag_1.name, tag_1.color)
        self.assertTrue(1 <= tag_1.color <= 11, "Default tag color must be a random integer between 1 and 11.")

        # 2. Test unique constraint on name
        _logger.info("Attempting to create duplicate Tag with name: %s", tag_1.name)
        with self.assertRaises(Exception, msg="Unique constraint name_uniq must reject duplicate tag name"):
            # Mute logger to avoid traceback prints for clean output
            with mute_logger('odoo.sql_db'):
                self.env['ir.attachment.tag'].create({
                    'name': 'Invoice Tag',
                })
                self.env.flush_all()
        _logger.info("Duplicate Tag creation correctly rejected with Unique Constraint exception.")

    def test_02_decode_content_xlsx(self):
        """Test decoding an Excel (.xlsx) file into HTML preview table."""
        _logger.info("========== RUNNING test_02_decode_content_xlsx ==========")
        # 1. Create in-memory pandas dataframe and convert to xlsx bytes
        df = pd.DataFrame({'Name': ['Alice', 'Bob'], 'Score': [95, 88]})
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
        xlsx_bytes = output.getvalue()
        xlsx_base64 = base64.b64encode(xlsx_bytes)
        _logger.info("In-memory XLSX document converted to base64.")

        # 2. Create attachment
        attachment = self.env['ir.attachment'].create({
            'name': 'scores.xlsx',
            'datas': xlsx_base64,
            'res_model': 'res.partner',
            'res_id': self.partner.id,
        })
        _logger.info("Created XLSX attachment ID: %s", attachment.id)

        # 3. Call decode_content
        html_table = self.env['ir.attachment'].decode_content(attachment.id, 'xlsx')
        _logger.info("XLSX content successfully decoded. HTML Table:\n%s", html_table)
        self.assertIn('<table', html_table, "XLSX decoding should produce an HTML table.")
        self.assertIn('Alice', html_table, "HTML preview should contain 'Alice'.")
        self.assertIn('Bob', html_table, "HTML preview should contain 'Bob'.")

    def test_03_decode_content_docx(self):
        """Test decoding a Word document (.docx) into paragraph list preview."""
        _logger.info("========== RUNNING test_03_decode_content_docx ==========")
        # 1. Create in-memory Word Docx
        doc = DocxDocument()
        doc.add_paragraph("Hello from Chatter Attachment Manager!")
        doc.add_paragraph("Odoo 19 integration test paragraph.")
        output = io.BytesIO()
        doc.save(output)
        docx_bytes = output.getvalue()
        docx_base64 = base64.b64encode(docx_bytes)
        _logger.info("In-memory DOCX document converted to base64.")

        # 2. Create attachment
        attachment = self.env['ir.attachment'].create({
            'name': 'document.docx',
            'datas': docx_base64,
            'res_model': 'res.partner',
            'res_id': self.partner.id,
        })
        _logger.info("Created DOCX attachment ID: %s", attachment.id)

        # 3. Call decode_content
        paragraphs = self.env['ir.attachment'].decode_content(attachment.id, 'docx')
        _logger.info("DOCX content successfully decoded. Paragraphs: %s", paragraphs)
        self.assertEqual(len(paragraphs), 2, "Should return list of 2 paragraphs.")
        self.assertEqual(paragraphs[0], "Hello from Chatter Attachment Manager!")
        self.assertEqual(paragraphs[1], "Odoo 19 integration test paragraph.")

    def test_04_save_edited_image(self):
        """Test replacing attachment file datas with edited image content."""
        _logger.info("========== RUNNING test_04_save_edited_image ==========")
        # 1. Create a dummy initial attachment
        attachment = self.env['ir.attachment'].create({
            'name': 'photo.png',
            'datas': base64.b64encode(b"initial_image_data"),
            'res_model': 'res.partner',
            'res_id': self.partner.id,
        })
        _logger.info("Created initial image attachment ID: %s", attachment.id)

        # 2. Call save_edited_image passing a base64 data URI
        updated_data_uri = "data:image/png;base64,dXBkYXRlZF9pbWFnZV9kYXRhX2Zyb21fdG9hc3Q="
        _logger.info("Calling save_edited_image with new base64 data URI.")
        self.env['ir.attachment'].save_edited_image(attachment.id, updated_data_uri)

        # 3. Check if attachment data matches the updated base64 string
        self.assertEqual(attachment.datas, base64.b64encode(b"updated_image_data_from_toast"))
        _logger.info("Image successfully updated on attachment. New data matches expected content.")

    def test_05_generate_qr_code(self):
        """Test dynamic generation of QR codes for downloading attachments."""
        _logger.info("========== RUNNING test_05_generate_qr_code ==========")
        # 1. Create base attachment
        attachment = self.env['ir.attachment'].create({
            'name': 'flyer.pdf',
            'datas': base64.b64encode(b"pdf_file_data"),
            'res_model': 'res.partner',
            'res_id': self.partner.id,
        })
        _logger.info("Created PDF attachment ID: %s", attachment.id)

        # 2. Generate QR code
        _logger.info("Generating download QR code...")
        qr_data = self.env['ir.attachment'].generate_qr_code(attachment.id)
        _logger.info("QR Code successfully generated: Company: %s, Image length: %s characters",
                     qr_data.get('company'), len(qr_data.get('image', '')))

        # 3. Assert return schema
        self.assertIn('image', qr_data, "QR code response must contain base64 image data.")
        self.assertIn('company', qr_data, "QR code response must contain current company name.")
        self.assertEqual(qr_data['company'], self.env.company.name)
        # Ensure image is valid base64
        decoded_qr = base64.b64decode(qr_data['image'])
        self.assertTrue(len(decoded_qr) > 0, "Decoded QR code must contain non-empty bytes.")
        _logger.info("========== TestAttachmentManager COMPLETED ALL TESTS SUCCESSFULLY ==========")
